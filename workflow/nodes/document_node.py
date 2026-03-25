"""文档生成节点：Word / 文本"""

import json, logging, os, time
from pathlib import Path
from workflow.registry import NodeRegistry, NodeTypeInfo
from workflow.types import WorkflowNode, NodeResult

logger = logging.getLogger(__name__)
_ROOT = Path(os.environ.get("AGENTFORGE_ROOT", ".")).resolve()


async def _document_executor(node: WorkflowNode, variables: dict, ctx: dict) -> NodeResult:
    action = node.config.get("action", "create_word")
    raw = node.config.get("path", "") or f"data/outputs/doc_{int(time.time())}.{'docx' if 'word' in action else 'txt'}"
    title = node.config.get("title", "")
    content = node.config.get("content", "")
    path = (_ROOT / raw).resolve()
    if not str(path).startswith(str(_ROOT)):
        return NodeResult(node_id=node.id, status="failed", error="路径越界")
    path.parent.mkdir(parents=True, exist_ok=True)

    if not content:
        last = ctx.get("_last_output", {})
        content = last.get("text", last.get("ai_result", "")) if isinstance(last, dict) else str(last)
        if isinstance(last, dict) and not content:
            content = json.dumps(last, ensure_ascii=False, indent=2)

    try:
        if action == "create_word":
            try:
                from docx import Document
                doc = Document()
                if title: doc.add_heading(title, level=1)
                for line in content.split("\n"):
                    l = line.strip()
                    if l.startswith("## "): doc.add_heading(l[3:], level=2)
                    elif l.startswith("# "): doc.add_heading(l[2:], level=1)
                    elif l: doc.add_paragraph(l)
                if not raw.endswith(".docx"): raw += ".docx"; path = (_ROOT / raw).resolve()
                doc.save(str(path))
            except ImportError:
                path.write_text(content, encoding="utf-8")
        else:
            if title: content = f"# {title}\n\n{content}"
            path.write_text(content, encoding="utf-8")
        return NodeResult(node_id=node.id, status="completed", output={"path": raw, "size": path.stat().st_size})
    except Exception as e:
        return NodeResult(node_id=node.id, status="failed", error=f"文档生成失败: {e}")


def register_document(registry: NodeRegistry) -> None:
    registry.register(NodeTypeInfo(name="document", display_name="文档生成", group="data", icon="file-text",
        description="生成 Word 或文本文档", parameters=[
            {"name": "action", "type": "options", "displayName": "操作", "default": "create_word",
             "options": [{"name": "Word", "value": "create_word"}, {"name": "文本", "value": "create_text"}]},
            {"name": "title", "type": "string", "displayName": "标题", "default": ""},
            {"name": "content", "type": "string", "displayName": "内容", "default": ""},
            {"name": "path", "type": "string", "displayName": "路径", "default": ""},
        ], executor=_document_executor))
