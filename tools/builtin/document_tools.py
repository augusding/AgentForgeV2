"""
AgentForge V2 — 文档工具集

excel_processor: Excel 读写 (openpyxl)
word_processor: Word 文档读写 (python-docx)
pdf_processor: PDF 读取/合并 (pdfplumber/PyPDF2)
"""

from __future__ import annotations

import json
import logging
import os
from pathlib import Path

from tools.registry import ToolDefinition

logger = logging.getLogger(__name__)

_ROOT = Path(os.environ.get("AGENTFORGE_ROOT", ".")).resolve()


def _safe_path(raw: str) -> Path:
    p = (_ROOT / raw).resolve()
    if not str(p).startswith(str(_ROOT)):
        raise PermissionError(f"路径越界: {raw}")
    return p


# ── Excel 处理 ────────────────────────────────────────────

async def _excel_handler(args: dict) -> str:
    try:
        import openpyxl
    except ImportError:
        return json.dumps({"error": "需要安装 openpyxl: pip install openpyxl"}, ensure_ascii=False)

    action = args.get("action", "")
    raw_path = args.get("path", "")
    if not raw_path:
        return json.dumps({"error": "缺少 path 参数"}, ensure_ascii=False)

    try:
        path = _safe_path(raw_path)
    except PermissionError as e:
        return json.dumps({"error": str(e)}, ensure_ascii=False)

    try:
        if action == "create":
            data_str = args.get("data", "")
            if not data_str:
                return json.dumps({"error": "create 需要 data 参数"}, ensure_ascii=False)
            rows = json.loads(data_str)
            wb = openpyxl.Workbook()
            ws = wb.active
            ws.title = args.get("sheet_name", "Sheet1")
            for row in rows:
                ws.append(row if isinstance(row, list) else [row])
            path.parent.mkdir(parents=True, exist_ok=True)
            wb.save(str(path))
            return json.dumps({"status": "created", "path": raw_path, "rows": len(rows)}, ensure_ascii=False)

        elif action == "read":
            if not path.is_file():
                return json.dumps({"error": f"文件不存在: {raw_path}"}, ensure_ascii=False)
            wb = openpyxl.load_workbook(str(path), read_only=True)
            sheet = args.get("sheet_name", "")
            ws = wb[sheet] if sheet and sheet in wb.sheetnames else wb.active
            result = []
            for row in ws.iter_rows(max_row=500, values_only=True):
                result.append([str(c) if c is not None else "" for c in row])
            wb.close()
            output = json.dumps(result, ensure_ascii=False)
            return output[:8000] if len(output) > 8000 else output

        elif action == "stats":
            if not path.is_file():
                return json.dumps({"error": f"文件不存在: {raw_path}"}, ensure_ascii=False)
            wb = openpyxl.load_workbook(str(path), read_only=True)
            info = {"sheets": []}
            for name in wb.sheetnames:
                ws = wb[name]
                info["sheets"].append({"name": name, "rows": ws.max_row or 0, "columns": ws.max_column or 0})
            wb.close()
            return json.dumps(info, ensure_ascii=False)

        elif action == "add_sheet":
            if not path.is_file():
                return json.dumps({"error": f"文件不存在: {raw_path}"}, ensure_ascii=False)
            wb = openpyxl.load_workbook(str(path))
            sheet_name = args.get("sheet_name", "NewSheet")
            ws = wb.create_sheet(title=sheet_name)
            data_str = args.get("data", "")
            if data_str:
                for row in json.loads(data_str):
                    ws.append(row if isinstance(row, list) else [row])
            wb.save(str(path))
            return json.dumps({"status": "sheet_added", "sheet": sheet_name}, ensure_ascii=False)

        return json.dumps({"error": f"未知操作: {action}"}, ensure_ascii=False)
    except Exception as e:
        return json.dumps({"error": f"Excel 操作失败: {e}"}, ensure_ascii=False)


excel_processor = ToolDefinition(
    name="excel_processor",
    description="Excel 表格处理。支持 create/read/stats/add_sheet 操作。",
    input_schema={
        "type": "object",
        "properties": {
            "action": {"type": "string", "enum": ["create", "read", "stats", "add_sheet"], "description": "操作类型"},
            "path": {"type": "string", "description": "文件路径(.xlsx)"},
            "data": {"type": "string", "description": "JSON 二维数组，如 [[\"A\",\"B\"],[1,2]]"},
            "sheet_name": {"type": "string", "description": "工作表名称"},
        },
        "required": ["action", "path"],
    },
    handler=_excel_handler,
    category="document",
)


# ── Word 处理 ─────────────────────────────────────────────

async def _word_handler(args: dict) -> str:
    try:
        from docx import Document
    except ImportError:
        return json.dumps({"error": "需要安装 python-docx: pip install python-docx"}, ensure_ascii=False)

    action = args.get("action", "")
    raw_path = args.get("path", "")
    if not raw_path:
        return json.dumps({"error": "缺少 path 参数"}, ensure_ascii=False)

    try:
        path = _safe_path(raw_path)
    except PermissionError as e:
        return json.dumps({"error": str(e)}, ensure_ascii=False)

    try:
        if action == "create":
            content = args.get("content", "")
            if not content:
                return json.dumps({"error": "create 需要 content 参数"}, ensure_ascii=False)
            doc = Document()
            for line in content.split("\n"):
                line = line.strip()
                if not line:
                    continue
                if line.startswith("## "):
                    doc.add_heading(line[3:], level=2)
                elif line.startswith("# "):
                    doc.add_heading(line[2:], level=1)
                else:
                    doc.add_paragraph(line)
            path.parent.mkdir(parents=True, exist_ok=True)
            doc.save(str(path))
            return json.dumps({"status": "created", "path": raw_path}, ensure_ascii=False)

        elif action == "read":
            if not path.is_file():
                return json.dumps({"error": f"文件不存在: {raw_path}"}, ensure_ascii=False)
            doc = Document(str(path))
            texts = []
            for para in doc.paragraphs[:200]:
                if para.text.strip():
                    texts.append(para.text)
            result = "\n".join(texts)
            return result[:8000] if len(result) > 8000 else result

        elif action == "replace":
            if not path.is_file():
                return json.dumps({"error": f"文件不存在: {raw_path}"}, ensure_ascii=False)
            find_text = args.get("find", "")
            replace_text = args.get("replace_with", "")
            doc = Document(str(path))
            count = 0
            for para in doc.paragraphs:
                if find_text in para.text:
                    para.text = para.text.replace(find_text, replace_text)
                    count += 1
            doc.save(str(path))
            return json.dumps({"status": "replaced", "count": count}, ensure_ascii=False)

        elif action == "add_table":
            if not path.is_file():
                return json.dumps({"error": f"文件不存在: {raw_path}"}, ensure_ascii=False)
            table_str = args.get("table_data", "")
            if not table_str:
                return json.dumps({"error": "add_table 需要 table_data"}, ensure_ascii=False)
            rows = json.loads(table_str)
            doc = Document(str(path))
            if rows:
                table = doc.add_table(rows=len(rows), cols=len(rows[0]))
                for i, row in enumerate(rows):
                    for j, cell in enumerate(row):
                        table.cell(i, j).text = str(cell)
            doc.save(str(path))
            return json.dumps({"status": "table_added", "rows": len(rows)}, ensure_ascii=False)

        return json.dumps({"error": f"未知操作: {action}"}, ensure_ascii=False)
    except Exception as e:
        return json.dumps({"error": f"Word 操作失败: {e}"}, ensure_ascii=False)


word_processor = ToolDefinition(
    name="word_processor",
    description="Word 文档处理。支持 create/read/replace/add_table 操作。",
    input_schema={
        "type": "object",
        "properties": {
            "action": {"type": "string", "enum": ["create", "read", "replace", "add_table"], "description": "操作类型"},
            "path": {"type": "string", "description": "文档路径(.docx)"},
            "content": {"type": "string", "description": "文档内容（create时，\\n分段，#开头为标题）"},
            "find": {"type": "string", "description": "查找文本（replace时）"},
            "replace_with": {"type": "string", "description": "替换文本（replace时）"},
            "table_data": {"type": "string", "description": "表格数据，JSON二维数组"},
        },
        "required": ["action", "path"],
    },
    handler=_word_handler,
    category="document",
)


# ── PDF 处理 ──────────────────────────────────────────────

async def _pdf_handler(args: dict) -> str:
    action = args.get("action", "")

    try:
        if action == "read":
            raw_path = args.get("path", "")
            if not raw_path:
                return json.dumps({"error": "缺少 path 参数"}, ensure_ascii=False)
            path = _safe_path(raw_path)
            if not path.is_file():
                return json.dumps({"error": f"文件不存在: {raw_path}"}, ensure_ascii=False)
            return _read_pdf(path)

        elif action == "info":
            raw_path = args.get("path", "")
            path = _safe_path(raw_path)
            if not path.is_file():
                return json.dumps({"error": f"文件不存在: {raw_path}"}, ensure_ascii=False)
            return _pdf_info(path)

        elif action == "merge":
            paths_str = args.get("paths", "[]")
            raw_paths = json.loads(paths_str)
            output = args.get("output_path", "data/outputs/merged.pdf")
            return await _merge_pdfs(raw_paths, output)

        return json.dumps({"error": f"未知操作: {action}"}, ensure_ascii=False)
    except Exception as e:
        return json.dumps({"error": f"PDF 操作失败: {e}"}, ensure_ascii=False)


def _read_pdf(path: Path) -> str:
    # 尝试 pdfplumber
    try:
        import pdfplumber
        texts = []
        with pdfplumber.open(str(path)) as pdf:
            for i, page in enumerate(pdf.pages[:50], 1):
                t = page.extract_text() or ""
                if t.strip():
                    texts.append(f"[Page {i}]\n{t}")
        result = "\n\n".join(texts)
        return result[:8000] if len(result) > 8000 else (result or "(无可提取文本)")
    except ImportError:
        pass

    # fallback: PyPDF2
    try:
        from PyPDF2 import PdfReader
        reader = PdfReader(str(path))
        texts = []
        for i, page in enumerate(reader.pages[:50], 1):
            t = page.extract_text() or ""
            if t.strip():
                texts.append(f"[Page {i}]\n{t}")
        result = "\n\n".join(texts)
        return result[:8000] if len(result) > 8000 else (result or "(无可提取文本)")
    except ImportError:
        return json.dumps({"error": "需要安装 pdfplumber 或 PyPDF2"}, ensure_ascii=False)


def _pdf_info(path: Path) -> str:
    try:
        from PyPDF2 import PdfReader
        reader = PdfReader(str(path))
        info = reader.metadata
        return json.dumps({
            "pages": len(reader.pages),
            "title": info.title if info else "",
            "author": info.author if info else "",
        }, ensure_ascii=False)
    except ImportError:
        return json.dumps({"error": "需要安装 PyPDF2"}, ensure_ascii=False)


async def _merge_pdfs(raw_paths: list[str], output_raw: str) -> str:
    try:
        from PyPDF2 import PdfWriter, PdfReader
    except ImportError:
        return json.dumps({"error": "需要安装 PyPDF2"}, ensure_ascii=False)

    writer = PdfWriter()
    for rp in raw_paths:
        path = _safe_path(rp)
        if not path.is_file():
            return json.dumps({"error": f"文件不存在: {rp}"}, ensure_ascii=False)
        reader = PdfReader(str(path))
        for page in reader.pages:
            writer.add_page(page)

    output = _safe_path(output_raw)
    output.parent.mkdir(parents=True, exist_ok=True)
    with open(str(output), "wb") as f:
        writer.write(f)
    return json.dumps({"status": "merged", "output": output_raw, "pages": len(writer.pages)}, ensure_ascii=False)


pdf_processor = ToolDefinition(
    name="pdf_processor",
    description="PDF 文件处理。支持 read/info/merge 操作。",
    input_schema={
        "type": "object",
        "properties": {
            "action": {"type": "string", "enum": ["read", "info", "merge"], "description": "操作类型"},
            "path": {"type": "string", "description": "PDF 文件路径"},
            "paths": {"type": "string", "description": "JSON 数组，多个 PDF 路径（merge时）"},
            "output_path": {"type": "string", "description": "输出路径（merge时）"},
        },
        "required": ["action"],
    },
    handler=_pdf_handler,
    category="document",
)


ALL_DOCUMENT_TOOLS = [excel_processor, word_processor, pdf_processor]
